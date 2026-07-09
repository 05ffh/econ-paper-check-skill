#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""经管论文智检 Skill - DOCX 报告渲染脚本（纯 I/O，无判断逻辑）。

用法：
    python scripts/render_report.py diagnostic_result.json --out 报告.docx --source 原论文.docx

只读取 diagnostic_result.json 字段，渲染为红黄绿灰分级、11 章节的正式 docx。
"""
from __future__ import annotations

import argparse
import datetime as _dt
import json
import sys
from pathlib import Path
from typing import Dict, List, Optional

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Pt, Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError as exc:  # pragma: no cover
    raise SystemExit("缺少 python-docx，请先执行：pip install python-docx") from exc

DISCLAIMER = (
    "本报告由“经管论文智检 Skill”基于用户上传的 Word 文档自动生成，仅用于本科论文和课程论文写作规范自检，"
    "不替代导师、答辩委员或学校正式评审意见。对于无法可靠解析的表格、公式、图片或主观方法选择，报告中已标注为"
    "“需人工确认”。本 Skill 不进行查重，不判断数据真实性，不联网核验参考文献真伪，也不提供论文代写服务。"
)



def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, rows: List[List[str]], header: bool = True) -> None:
    if not rows:
        return
    table = document.add_table(rows=0, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx, value in enumerate(row):
            set_cell_text(cells[c_idx], value, bold=(header and r_idx == 0))
            if header and r_idx == 0:
                set_cell_shading(cells[c_idx], "D9EAF7")
    document.add_paragraph("")


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_para(document: Document, text: str, bold_label: Optional[str] = None) -> None:
    p = document.add_paragraph()
    if bold_label:
        r = p.add_run(bold_label)
        r.bold = True
        r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)


def _issue_card(doc: Document, idx: int, f: dict) -> None:
    add_heading(doc, f"{idx}. {f.get('issue_id', '')}｜{f.get('issue_type', '')}", 2)
    rows = [
        ["项目", "内容"],
        ["问题等级", f.get("level", "")],
        ["所属诊断域", f.get("domain", "")],
        ["所在位置", f.get("location", "")],
        ["原文证据", f.get("evidence", "")],
        ["证据强度", f.get("evidence_strength", "")],
        ["问题说明", f.get("explanation", "")],
        ["修改建议", f.get("suggestion", "")],
        ["需人工确认", "是" if f.get("need_manual_confirmation") else "否"],
    ]
    add_table(doc, rows)


def write_report(result: Dict[str, object], source_name: str, output_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    profile = result.get("paper_profile", {}) or {}
    counts = result.get("summary_counts", {}) or {}
    today = _dt.datetime.now().strftime("%Y-%m-%d")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("经管本科论文智检报告")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("基于经管论文智检 Skill 的提交前自检反馈").font.size = Pt(12)
    doc.add_paragraph("")

    add_heading(doc, "一、检测基本信息", 1)
    add_table(doc, [
        ["项目", "内容"],
        ["原论文文件名", source_name],
        ["检测日期", today],
        ["论文类型", str(profile.get("paper_type", "不确定，需人工确认"))],
        ["数据类型", str(profile.get("data_type", "不确定，需人工确认"))],
        ["已识别方法", "、".join(profile.get("detected_methods", []) or ["未识别"])],
        ["总体风险等级", str(result.get("overall_risk_level", "需人工确认"))],
        ["报告生成说明", "本报告基于 DOCX 可解析文本与表格自动生成。"],
    ])

    add_heading(doc, "二、论文画像识别结果", 1)
    tp = profile.get("trigger_plan", {}) or {}
    add_table(doc, [
        ["画像字段", "识别结果"],
        ["论文类型", str(profile.get("paper_type", "不确定，需人工确认"))],
        ["数据类型", str(profile.get("data_type", "不确定，需人工确认"))],
        ["样本对象", str(profile.get("sample_object", "不确定，需人工确认"))],
        ["样本期间", str(profile.get("sample_period", "不确定，需人工确认"))],
        ["已识别方法", "、".join(profile.get("detected_methods", []) or ["未识别"])],
        ["触发规则组", "、".join(tp.get("triggered_rule_groups", []) or ["无"])],
        ["不适用规则组", "、".join(tp.get("not_applicable_rule_groups", []) or ["无"])],
        ["需人工确认事项", "、".join(tp.get("manual_confirmation_items", []) or ["无"])],
    ])

    add_heading(doc, "三、总体评价与修改优先级", 1)
    add_para(doc, f"综合当前可识别内容，本报告将论文提交前修改风险判断为：{result.get('overall_risk_level', '需人工确认')}。")
    add_table(doc, [
        ["类别", "数量"],
        ["通过项", str(counts.get("pass", 0))],
        ["红色必改问题", str(counts.get("red", 0))],
        ["黄色建议补充", str(counts.get("yellow", 0))],
        ["绿色优化提升", str(counts.get("green", 0))],
        ["灰色需人工确认", str(counts.get("manual", 0))],
        ["不适用项", str(counts.get("na", 0))],
    ])

    add_heading(doc, "四、通过项概览", 1)
    passes = result.get("pass_items", []) or []
    if passes:
        rows = [["编号", "所属诊断域", "通过项", "证据"]]
        for p in passes:
            rows.append([p.get("id", ""), p.get("domain", ""), p.get("text", ""), p.get("evidence", "")])
        add_table(doc, rows)
    else:
        add_para(doc, "本次检测未单独列出通过项。")

    issues = result.get("issues", []) or []
    by_level = {lv: [f for f in issues if f.get("level") == lv] for lv in ["红色", "黄色", "绿色", "灰色"]}
    manual_items = result.get("manual_confirmation_items", []) or []

    add_heading(doc, "五、红色必改问题", 1)
    if by_level["红色"]:
        for i, f in enumerate(by_level["红色"], 1):
            _issue_card(doc, i, f)
    else:
        add_para(doc, "本次检测未发现明确的红色必改问题。仍建议结合导师意见进一步复核论文核心内容。")

    add_heading(doc, "六、黄色建议补充问题", 1)
    if by_level["黄色"]:
        for i, f in enumerate(by_level["黄色"], 1):
            _issue_card(doc, i, f)
    else:
        add_para(doc, "本次检测未发现明显需要补充的黄色问题。")

    add_heading(doc, "七、绿色优化提升建议", 1)
    if by_level["绿色"]:
        for i, f in enumerate(by_level["绿色"], 1):
            _issue_card(doc, i, f)
    else:
        add_para(doc, "本次检测暂未生成绿色优化提升建议。")

    add_heading(doc, "八、灰色需人工确认事项", 1)
    gray = by_level["灰色"]
    if gray or manual_items:
        for i, f in enumerate(gray, 1):
            _issue_card(doc, i, f)
        if manual_items:
            rows = [["编号", "确认事项", "所在位置", "需要确认的原因", "建议人工复核方式"]]
            for m in manual_items:
                rows.append([m.get("id", ""), m.get("item", ""), m.get("location", ""),
                             m.get("reason", ""), m.get("how_to_check", "")])
            add_table(doc, rows)
    else:
        add_para(doc, "本次检测未发现必须人工确认的解析事项。")

    add_heading(doc, "九、五大诊断域详情", 1)
    domains = ["选题与研究问题质检", "文献综述与理论链条质检", "数据、变量与口径质检",
               "方法模型与实证结果质检", "结构表达与学术规范质检"]
    for d in domains:
        add_heading(doc, d, 2)
        d_issues = [f.get("issue_id", "") for f in issues if f.get("domain") == d]
        d_pass = [p.get("id", "") for p in passes if p.get("domain") == d]
        add_para(doc, "通过项编号：" + ("、".join(d_pass) if d_pass else "无"))
        add_para(doc, "问题编号：" + ("、".join(d_issues) if d_issues else "无"))

    add_heading(doc, "十、优先修改行动清单", 1)
    actions = result.get("priority_actions", []) or []
    if actions:
        rows = [["优先级", "修改行动", "对应问题编号", "预期效果"]]
        for a in actions:
            rows.append([str(a.get("priority", "")), a.get("action", ""),
                         a.get("issue_ref", ""), a.get("expected", "")])
        add_table(doc, rows)
    else:
        add_para(doc, "本次检测未生成优先修改行动清单。")

    add_heading(doc, "十一、免责声明", 1)
    add_para(doc, DISCLAIMER)

    na_items = result.get("not_applicable_items", []) or []
    if na_items:
        add_heading(doc, "附：不适用规则说明", 2)
        rows = [["编号", "规则组", "不适用原因"]]
        for n in na_items:
            rows.append([n.get("id", ""), n.get("rule_group", ""), n.get("reason", "")])
        add_table(doc, rows)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(description="将 diagnostic_result.json 渲染为 DOCX 报告")
    parser.add_argument("result", help="diagnostic_result.json 路径")
    parser.add_argument("--out", "-o", help="输出 DOCX 报告路径")
    parser.add_argument("--source", help="原论文文件名（用于报告命名与基本信息）", default=None)
    args = parser.parse_args(argv)

    result_path = Path(args.result).expanduser().resolve()
    if not result_path.exists():
        print(f"错误：结果文件不存在：{result_path}", file=sys.stderr)
        return 2
    result = json.loads(result_path.read_text(encoding="utf-8"))

    source_name = Path(args.source).name if args.source else "未提供"
    today = _dt.datetime.now().strftime("%Y%m%d")
    if args.out:
        output_path = Path(args.out).expanduser().resolve()
    else:
        stem = Path(source_name).stem if source_name != "未提供" else "论文"
        output_path = result_path.with_name(f"经管论文智检报告_{stem}_{today}.docx")

    write_report(result, source_name, output_path)
    print(f"已生成 DOCX 质检报告：{output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
