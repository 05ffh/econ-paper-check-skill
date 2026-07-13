#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""self_check.py · 用户侧轻量安装自检

**目的**：让接收 Skill 的普通用户确认当前安装是否可正常运行，
         **不涉及 Benchmark、不调 Ark、不碰任何私有样本**。

**硬约束**（见 plans/M4_BUILD_GATE_ADDENDUM.md §8）：
  - 默认不调用真实 Ark
  - 默认不使用学生论文和私有 Benchmark 样本
  - 使用内置合成数据或最小公开 fixture
  - 自检只回答"当前安装能否运行"，不宣称"质检准确率已重新验证"
  - 完整 Benchmark 仍只供开发者和发布流程使用

**依赖**：只依赖 `requirements-core.txt`。

CLI:
    python scripts/self_check.py                # 基础自检（默认）
    python scripts/self_check.py --json         # 机器可读
    python scripts/self_check.py --with-vision  # 可选：合成图片走一次 Ark
"""
from __future__ import annotations
import argparse
import importlib
import io
import json
import sys
from pathlib import Path

_REPO_ROOT = Path(__file__).resolve().parent.parent


# ---------- 分项检查 ----------

def _check_import(module: str, human_name: str) -> tuple[bool, str]:
    try:
        importlib.import_module(module)
        return True, f"✅ {human_name}"
    except ImportError as e:
        return False, f"❌ {human_name}（未安装：{e}）"


def check_docx_parse() -> tuple[bool, str]:
    """用内置合成 docx 走一次 python-docx 读取。"""
    try:
        from docx import Document
        buf = io.BytesIO()
        doc = Document()
        doc.add_paragraph("self_check synthetic paragraph 1")
        doc.add_paragraph("self_check synthetic paragraph 2")
        doc.save(buf)
        buf.seek(0)
        doc2 = Document(buf)
        paras = [p.text for p in doc2.paragraphs if p.text.strip()]
        if len(paras) >= 2:
            return True, "✅ DOCX 基础解析（合成样本 2 段）"
        return False, f"❌ DOCX 解析返回段数异常：{len(paras)}"
    except Exception as e:
        return False, f"❌ DOCX 解析异常：{e}"


def check_pdf_parse() -> tuple[bool, str]:
    """检查 pypdf 与 pdfplumber 是否可 import 并能处理最小 PDF。"""
    try:
        import pypdf  # noqa
        import pdfplumber  # noqa
        # 只做 import 级别校验；避免依赖外部 fixture
        return True, "✅ 文本型 PDF 解析（pypdf + pdfplumber 可用）"
    except Exception as e:
        return False, f"❌ PDF 解析依赖缺失：{e}"


def check_kb_a() -> tuple[bool, str]:
    """检查 KB-A 规范层能否加载（YAML 文件计数）。"""
    kb_a = _REPO_ROOT / "knowledge_base" / "norms"
    if not kb_a.exists():
        return False, "❌ KB-A 目录不存在：knowledge_base/norms/"
    yamls = list(kb_a.rglob("*.yaml"))
    total_rules = 0
    for p in yamls:
        try:
            import yaml
            with open(p, "r", encoding="utf-8") as f:
                data = yaml.safe_load(f) or {}
            rules = data.get("norms") or data.get("rules") or []
            total_rules += len(rules)
        except Exception:
            continue
    if total_rules >= 1:
        return True, f"✅ KB-A 规范加载（{len(yamls)} 文件 · {total_rules} 条）"
    return False, "❌ KB-A 未找到有效规则"


def check_report_render() -> tuple[bool, str]:
    """检查 render_report_html 能生成非空 HTML。"""
    try:
        # 只 import 不实跑，避免耦合到 rules 逻辑
        import jinja2  # noqa
        return True, "✅ 报告渲染引擎（jinja2 可用）"
    except Exception as e:
        return False, f"❌ 报告渲染依赖缺失：{e}"


def check_vision_graceful_degrade() -> tuple[bool, str]:
    """未配置视觉时，vision_pipeline 应可 import 且能 fallback。"""
    try:
        # vision 子模块可能未装 openai，允许 ImportError 走降级路径
        from scripts import vision_hint  # noqa
        return True, "⚪ 视觉未配置（降级路径可用；用户可选装 requirements-vision.txt）"
    except ImportError:
        return True, "⚪ 视觉子系统未加载（降级路径生效）"
    except Exception as e:
        return False, f"❌ 视觉降级路径异常：{e}"


def check_ark_smoke(fixture: Path) -> tuple[bool, str]:
    """可选：用合成 PNG 走一次 Ark。仅 --with-vision 时执行。"""
    if not fixture.exists():
        return False, f"❌ 视觉冒烟 fixture 不存在：{fixture}"
    try:
        from scripts.vision.providers.ark_provider import ArkVisionProvider  # noqa
    except Exception as e:
        return False, f"❌ 视觉未装：{e}"
    # 依赖 .env.local 中的 ARK_* 变量；此处不透露细节
    try:
        import os
        if not os.environ.get("ARK_API_KEY"):
            return False, "❌ ARK_API_KEY 未配置，跳过。"
        # 真调一次
        from scripts.vision.providers.ark_provider import ArkVisionProvider
        prov = ArkVisionProvider()
        _ = prov.recognize(fixture)
        return True, "✅ 视觉冒烟通过（Ark 真调 1 次，合成 PNG）"
    except Exception as e:
        return False, f"❌ 视觉冒烟失败：{e}"


# ---------- 主逻辑 ----------

def run(with_vision: bool = False) -> dict:
    results = []

    for m, name in [
        ("docx", "python-docx"),
        ("pypdf", "pypdf"),
        ("pdfplumber", "pdfplumber"),
        ("yaml", "pyyaml"),
        ("jinja2", "jinja2"),
    ]:
        ok, msg = _check_import(m, name)
        results.append({"section": "deps", "ok": ok, "msg": msg})

    for fn, section in [
        (check_docx_parse, "docx"),
        (check_pdf_parse, "pdf"),
        (check_kb_a, "kb_a"),
        (check_report_render, "render"),
        (check_vision_graceful_degrade, "vision_degrade"),
    ]:
        ok, msg = fn()
        results.append({"section": section, "ok": ok, "msg": msg})

    if with_vision:
        fixture = _REPO_ROOT / "tests" / "vision" / "fixtures" / "smoke_synthetic_table.png"
        ok, msg = check_ark_smoke(fixture)
        results.append({"section": "vision_smoke", "ok": ok, "msg": msg})

    return {
        "ok": all(r["ok"] for r in results),
        "results": results,
        "disclaimer": (
            "self_check 只回答“当前安装能否运行”，不代表质检准确率已重新验证。"
            "完整 Benchmark 由开发者运行，见 benchmarks/README.md。"
        ),
    }


def main() -> int:
    ap = argparse.ArgumentParser(description="用户侧轻量安装自检（不涉及 Benchmark，不调 Ark）")
    ap.add_argument("--json", action="store_true", help="输出 JSON")
    ap.add_argument("--with-vision", action="store_true",
                    help="可选：用合成 PNG 走一次 Ark（需 ARK_API_KEY）")
    args = ap.parse_args()

    r = run(with_vision=args.with_vision)

    if args.json:
        print(json.dumps(r, ensure_ascii=False, indent=2))
    else:
        print("=" * 60)
        print("经管论文智检 Skill · 安装自检")
        print("=" * 60)
        for item in r["results"]:
            print(f"  [{item['section']:15s}] {item['msg']}")
        print("=" * 60)
        print(f"总体：{'✅ 通过' if r['ok'] else '❌ 有失败项'}")
        print(f"\n⚠️  {r['disclaimer']}")

    return 0 if r["ok"] else 1


if __name__ == "__main__":
    sys.exit(main())
